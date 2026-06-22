from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri"
FONT_CODE = "Consolas"
COLOR_NAVY = RGBColor(31, 61, 90)
COLOR_BLUE = RGBColor(46, 116, 181)
COLOR_DARK = RGBColor(33, 37, 41)
COLOR_MUTED = RGBColor(90, 98, 104)
COLOR_LIGHT = "E8EEF5"
COLOR_NOTE = "F6F8FB"
COLOR_BORDER = "C8D3E2"


SOURCES = [
    "https://code.claude.com/docs/en/setup",
    "https://code.claude.com/docs/en/authentication",
    "https://code.claude.com/docs/en/troubleshoot-install",
    "https://code.claude.com/docs/en/terminal-guide",
    "https://code.claude.com/docs/en/settings",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    dxa = str(int(width_inches * 1440))
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), dxa)


def set_table_borders(table, color="DADCE0", size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_indent(table, dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, name=FONT_BODY, size=11, bold=False, color=COLOR_DARK, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_style(document, name, style_type, base=None):
    styles = document.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, style_type)
    if base:
        style.base_style = styles[base]
    return style


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    title.font.name = FONT_HEADING
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
    title.font.size = Pt(24)
    title.font.color.rgb = COLOR_NAVY

    heading1 = document.styles["Heading 1"]
    heading1.font.name = FONT_HEADING
    heading1._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = COLOR_BLUE
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)

    heading2 = document.styles["Heading 2"]
    heading2.font.name = FONT_HEADING
    heading2._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
    heading2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
    heading2.font.size = Pt(13)
    heading2.font.color.rgb = COLOR_BLUE
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(7)

    heading3 = document.styles["Heading 3"]
    heading3.font.name = FONT_HEADING
    heading3._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
    heading3._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
    heading3.font.size = Pt(12)
    heading3.font.color.rgb = COLOR_NAVY
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(5)

    small = add_style(document, "SmallText", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    small.font.size = Pt(9.5)
    small.font.color.rgb = COLOR_MUTED
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.15

    code = add_style(document, "CodeBlock", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.1)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.1

    tip = add_style(document, "TipBox", WD_STYLE_TYPE.PARAGRAPH, "Normal")
    tip.paragraph_format.left_indent = Inches(0.1)
    tip.paragraph_format.right_indent = Inches(0.1)
    tip.paragraph_format.space_before = Pt(4)
    tip.paragraph_format.space_after = Pt(6)
    tip.paragraph_format.line_spacing = 1.15

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Guida operativa Claude Code")
    set_run_font(run, size=9, color=COLOR_MUTED)


def add_divider(document):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.0)
    run = p.add_run(" ")
    set_run_font(run, size=1, color=COLOR_MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_BORDER)


def add_title_block(document, title_text, subtitle_text):
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run(title_text)
    set_run_font(run, name=FONT_HEADING, size=24, bold=True, color=COLOR_NAVY)

    subtitle = document.add_paragraph(style="SmallText")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(subtitle, before=0, after=8, line=1.15)
    run = subtitle.add_run(subtitle_text)
    set_run_font(run, size=9.5, color=COLOR_MUTED)
    add_divider(document)


def add_paragraph(document, text, style="Normal", bold_prefix=None):
    p = document.add_paragraph(style=style)
    set_paragraph_spacing(p, after=6, line=1.25 if style == "Normal" else 1.15)
    if bold_prefix and text.startswith(bold_prefix):
        prefix, rest = text.split(":", 1)
        run = p.add_run(prefix + ":")
        set_run_font(run, bold=True)
        run2 = p.add_run(rest)
        set_run_font(run2)
    else:
        run = p.add_run(text)
        size = 9.5 if style == "SmallText" else 11
        set_run_font(run, size=size, color=COLOR_MUTED if style == "SmallText" else COLOR_DARK)
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run)


def add_code_block(document, lines):
    p = document.add_paragraph(style="CodeBlock")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name=FONT_CODE, size=9.5, color=COLOR_DARK)
        if idx < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)
    p_pr = p._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), COLOR_NOTE)


def add_callout(document, label, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color=COLOR_BORDER, size="8")
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 6.3)
    set_cell_shading(cell, COLOR_LIGHT)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=4, after=4, line=1.15)
    run = p.add_run(f"{label}: ")
    set_run_font(run, bold=True, color=COLOR_NAVY)
    run2 = p.add_run(text)
    set_run_font(run2, color=COLOR_DARK)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color=COLOR_BORDER, size="8")
    set_table_indent(table, 120)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_width(header_cells[idx], widths[idx])
        set_cell_shading(header_cells[idx], COLOR_LIGHT)
        p = header_cells[idx].paragraphs[0]
        set_paragraph_spacing(p, before=2, after=2, line=1.0)
        run = p.add_run(header)
        set_run_font(run, bold=True, color=COLOR_NAVY)
        header_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(row_cells[idx], widths[idx])
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row_cells[idx].paragraphs[0]
            set_paragraph_spacing(p, before=2, after=2, line=1.15)
            parts = value.split("\n")
            for part_idx, part in enumerate(parts):
                run = p.add_run(part)
                set_run_font(run, size=10.5)
                if part_idx < len(parts) - 1:
                    run.add_break(WD_BREAK.LINE)
    document.add_paragraph("")


def add_sources_section(document):
    document.add_heading("Fonti ufficiali utilizzate", level=1)
    for source in SOURCES:
        p = document.add_paragraph(style="SmallText")
        set_paragraph_spacing(p, after=2, line=1.05)
        run = p.add_run(source)
        set_run_font(run, size=9.2, color=COLOR_MUTED)


def build_standard_guide():
    doc = Document()
    configure_document(doc)
    add_title_block(
        doc,
        "Guida installazione Claude Code su Windows",
        "Documento operativo in italiano per utenti finali. Percorso consigliato: Windows nativo + PowerShell.",
    )

    add_paragraph(
        doc,
        "Questa guida aiuta a installare Claude Code su PC Windows, verificare che il comando funzioni e completare il primo accesso con il flusso standard via browser. Se il tuo team lavora con una chiave API Anthropic invece del login classico, usa la guida separata dedicata all'API key.",
    )
    add_callout(
        doc,
        "Scelta consigliata",
        "Per la maggior parte degli utenti Windows conviene usare l'installazione nativa con PowerShell. WSL ha senso solo se i tuoi progetti e toolchain vivono gia in ambiente Linux.",
    )

    doc.add_heading("Requisiti minimi", level=1)
    add_bullets(
        doc,
        [
            "Windows 10 versione 1809 o successiva, oppure Windows Server 2019 o successivo.",
            "Almeno 4 GB di RAM e processore x64 oppure ARM64.",
            "Connessione Internet funzionante e accesso a un paese supportato da Anthropic.",
            "Una shell disponibile: PowerShell, CMD, Bash o Zsh. Su Windows, PowerShell e sufficiente per partire.",
            "Per l'accesso standard servono credenziali Claude compatibili: Pro, Max, Team, Enterprise oppure Console. Il piano gratuito di Claude.ai non include Claude Code.",
        ],
    )

    doc.add_heading("PowerShell, CMD e WSL in breve", level=1)
    add_table(
        doc,
        ["Opzione", "Quando usarla", "Note pratiche"],
        [
            [
                "PowerShell",
                "Percorso principale consigliato.",
                "E preinstallato su Windows e usa il comando ufficiale `irm ... | iex`.",
            ],
            [
                "CMD",
                "Alternativa se lavori gia da Prompt dei comandi.",
                "Usa un comando diverso. Se lanci il comando CMD dentro PowerShell compare l'errore su `&&`.",
            ],
            [
                "WSL 2",
                "Solo se lavori su progetti Linux o ti serve il sandboxing.",
                "Si installa e si usa dentro la distribuzione WSL, non da PowerShell. WSL 1 non supporta il percorso migliore.",
            ],
        ],
        [1.1, 2.0, 3.2],
    )

    doc.add_heading("Git for Windows: opzionale ma utile", level=1)
    add_paragraph(
        doc,
        "Git for Windows non e obbligatorio per installare Claude Code, ma aggiunge Git Bash. Senza Git for Windows, Claude Code esegue i comandi shell con lo strumento PowerShell. Con Git for Windows puo usare anche il Bash tool, comodo se il progetto contiene script Bash o istruzioni pensate per ambienti Unix.",
    )
    add_callout(
        doc,
        "Quando installarlo",
        "Se non sai se ti serve, puoi iniziare senza Git for Windows. Se in seguito lavori con script Bash o repository che lo richiedono, puoi aggiungerlo senza reinstallare Claude Code.",
    )

    doc.add_heading("Procedura di installazione consigliata", level=1)
    doc.add_heading("1. Apri il terminale corretto", level=2)
    add_numbered(
        doc,
        [
            "Premi `Win + X` e apri Windows PowerShell oppure Windows Terminal.",
            "Controlla il prompt: in PowerShell vedi `PS C:\\Users\\NomeUtente>`. In CMD il prefisso `PS` non compare.",
            "Se vuoi usare Git Bash in futuro, puoi installare Git for Windows anche prima di installare Claude Code.",
        ],
    )

    doc.add_heading("2. Esegui il comando ufficiale in PowerShell", level=2)
    add_code_block(doc, ["irm https://claude.ai/install.ps1 | iex"])
    add_paragraph(
        doc,
        "Non serve aprire PowerShell come amministratore. Al termine dell'installazione potrai lanciare `claude` da qualsiasi nuova finestra di terminale.",
    )

    doc.add_heading("3. Alternative utili", level=2)
    add_paragraph(doc, "Se non vuoi usare PowerShell, hai due strade alternative:")
    add_paragraph(doc, "Alternativa CMD", style="SmallText")
    add_code_block(doc, ["curl -fsSL https://claude.ai/install.cmd -o install.cmd && install.cmd && del install.cmd"])
    add_paragraph(doc, "Alternativa WinGet", style="SmallText")
    add_code_block(doc, ["winget install Anthropic.ClaudeCode"])
    add_callout(
        doc,
        "Piano B consigliato",
        "WinGet e utile quando il download via script incontra restrizioni di rete o policy aziendali. Ricorda pero che gli aggiornamenti WinGet non sono automatici.",
    )

    doc.add_heading("Verifica post-installazione", level=1)
    add_paragraph(doc, "Dopo l'installazione chiudi e riapri il terminale, poi verifica il comando.")
    add_code_block(doc, ["claude --version", "claude doctor"])
    add_bullets(
        doc,
        [
            "`claude --version` conferma che il binario e raggiungibile dal PATH.",
            "`claude doctor` esegue un controllo piu ampio su installazione, aggiornamenti e configurazione.",
        ],
    )

    doc.add_heading("Primo avvio e login standard", level=1)
    add_numbered(
        doc,
        [
            "Digita `claude` nel terminale.",
            "Al primo avvio si apre il browser per completare il login.",
            "Se il browser non parte, premi `c` per copiare l'URL di accesso e incollalo manualmente nel browser.",
            "Se dopo il login il browser mostra un codice invece di ritornare al terminale, copia quel codice e incollalo nel prompt `Paste code here if prompted`.",
            "Se lavori normalmente con API key invece che con login browser, interrompi qui e segui la guida dedicata all'API key.",
        ],
    )
    add_callout(
        doc,
        "Nota importante",
        "Il piano gratuito Claude.ai non abilita Claude Code. Per il login standard servono Pro, Max, Team, Enterprise oppure credenziali Console assegnate dal tuo team.",
    )

    doc.add_heading("Errori frequenti e soluzione rapida", level=1)
    add_table(
        doc,
        ["Problema", "Causa probabile", "Cosa fare"],
        [
            [
                "`irm` non riconosciuto",
                "Hai aperto CMD invece di PowerShell.",
                "Apri PowerShell e rilancia `irm https://claude.ai/install.ps1 | iex`.",
            ],
            [
                "`&&` non valido",
                "Hai eseguito il comando CMD dentro PowerShell.",
                "Resta in PowerShell e usa il comando `irm ... | iex`, oppure apri davvero CMD.",
            ],
            [
                "`bash` non riconosciuto",
                "Hai copiato il comando macOS/Linux su Windows.",
                "Su Windows usa il comando PowerShell ufficiale.",
            ],
            [
                "`claude` non riconosciuto dopo l'installazione",
                "Il PATH non e stato ricaricato.",
                "Chiudi il terminale, riaprilo e riprova. Se persiste, controlla che l'installazione sia andata a buon fine con `claude doctor`.",
            ],
            [
                "Si apre Claude Desktop invece della CLI",
                "Una vecchia installazione di Claude Desktop prende priorita sul comando `claude`.",
                "Aggiorna Claude Desktop alla versione piu recente.",
            ],
            [
                "`The process cannot access the file` durante il download",
                "Un vecchio tentativo o l'antivirus sta bloccando `%USERPROFILE%\\.claude\\downloads`.",
                "Chiudi gli altri terminali, elimina la cartella download con `Remove-Item -Recurse -Force \"$env:USERPROFILE\\.claude\\downloads\"` e rilancia l'installer.",
            ],
            [
                "Errore: serve Git for Windows o PowerShell",
                "Claude Code non trova nessuna shell disponibile.",
                "Verifica che PowerShell sia nel PATH, oppure installa Git for Windows e riapri il terminale.",
            ],
            [
                "Errore legato a PowerShell x86 o Windows 32 bit",
                "Stai usando la versione 32 bit della shell oppure un ambiente non supportato.",
                "Apri Windows PowerShell standard o Windows Terminal. In caso di dubbio, esegui `[Environment]::Is64BitOperatingSystem`.",
            ],
        ],
        [1.8, 1.8, 2.6],
    )

    doc.add_heading("Best practice finali", level=1)
    add_bullets(
        doc,
        [
            "Non usare privilegi amministrativi se non richiesti dalla tua policy aziendale.",
            "Dopo installazioni o modifiche al PATH, chiudi e riapri il terminale prima di testare `claude`.",
            "Se sei alle prime armi, resta su PowerShell: e il percorso piu lineare e meglio documentato.",
            "Se installi Git for Windows, mantieni l'opzione consigliata per aggiungerlo al PATH.",
            "Se la rete aziendale blocca il download, prova WinGet oppure confrontati con chi gestisce proxy e firewall.",
        ],
    )

    doc.add_heading("Riepilogo operativo", level=1)
    add_numbered(
        doc,
        [
            "Apri PowerShell.",
            "Esegui `irm https://claude.ai/install.ps1 | iex`.",
            "Riapri il terminale e verifica con `claude --version` e `claude doctor`.",
            "Avvia `claude` e completa il login nel browser.",
        ],
    )

    add_sources_section(doc)
    return doc


def build_api_key_guide():
    doc = Document()
    configure_document(doc)
    add_title_block(
        doc,
        "Guida Claude Code su Windows con API key",
        "Documento operativo in italiano per utenti finali che usano Claude Code con `ANTHROPIC_API_KEY` invece del login browser basato su subscription.",
    )

    add_paragraph(
        doc,
        "Questa guida spiega come installare Claude Code su Windows e farlo usare con una API key Anthropic. Il focus non e l'amministrazione della piattaforma, ma il percorso pratico per un utente che ha gia ricevuto una chiave valida dal proprio team.",
    )
    add_callout(
        doc,
        "Differenza chiave",
        "Nel flusso API key non ti affidi al login standard via browser come credenziale primaria. Claude Code puo leggere `ANTHROPIC_API_KEY` direttamente dall'ambiente del terminale.",
    )

    doc.add_heading("Prerequisiti", level=1)
    add_bullets(
        doc,
        [
            "Windows 10 versione 1809 o successiva, oppure Windows Server 2019 o successivo.",
            "Claude Code installato oppure installabile con il percorso Windows nativo + PowerShell.",
            "Una API key valida gia rilasciata dal tuo team o dalla tua organizzazione.",
            "Se la chiave arriva da Anthropic Console, il tuo account deve avere il ruolo adeguato: `Claude Code` oppure `Developer`.",
        ],
    )

    doc.add_heading("Installazione Windows in sintesi", level=1)
    add_paragraph(
        doc,
        "Se Claude Code non e ancora installato, il percorso consigliato resta PowerShell.",
    )
    add_code_block(doc, ["irm https://claude.ai/install.ps1 | iex"])
    add_paragraph(doc, "Alternative utili se necessario:", style="SmallText")
    add_code_block(
        doc,
        [
            "curl -fsSL https://claude.ai/install.cmd -o install.cmd && install.cmd && del install.cmd",
            "winget install Anthropic.ClaudeCode",
        ],
    )
    add_paragraph(
        doc,
        "Dopo l'installazione verifica sempre il binario con `claude --version` e `claude doctor`.",
    )

    doc.add_heading("Configurare la API key in PowerShell", level=1)
    doc.add_heading("1. Variabile valida solo per la sessione corrente", level=2)
    add_code_block(doc, ['$env:ANTHROPIC_API_KEY="inserisci-qui-la-tua-api-key"'])
    add_paragraph(
        doc,
        "Questa opzione e la piu sicura per test veloci: la chiave sparisce quando chiudi la finestra PowerShell.",
    )

    doc.add_heading("2. Variabile persistente a livello utente", level=2)
    add_code_block(
        doc,
        ['[System.Environment]::SetEnvironmentVariable("ANTHROPIC_API_KEY","inserisci-qui-la-tua-api-key","User")']
    )
    add_paragraph(
        doc,
        "Dopo averla impostata, chiudi e riapri PowerShell. La variabile sara disponibile per le nuove sessioni del tuo utente Windows.",
    )
    add_callout(
        doc,
        "Sicurezza",
        "Evita di salvare la chiave dentro repository, script condivisi o screenshot. Se possibile preferisci la variabile di sessione per attivita temporanee.",
    )

    doc.add_heading("3. Verifica senza stampare la chiave completa", level=2)
    add_code_block(
        doc,
        [
            'if ($env:ANTHROPIC_API_KEY) { "ANTHROPIC_API_KEY configurata" } else { "ANTHROPIC_API_KEY assente" }',
            '[Environment]::GetEnvironmentVariable("ANTHROPIC_API_KEY","User")',
        ],
    )
    add_paragraph(
        doc,
        "Il primo comando controlla la chiave nella sessione corrente. Il secondo verifica se la variabile e stata salvata in modo persistente per l'utente.",
    )

    doc.add_heading("Come Claude Code sceglie le credenziali", level=1)
    add_paragraph(
        doc,
        "Quando sono presenti piu credenziali, Claude Code segue un ordine di precedenza. Per il percorso che interessa qui, `ANTHROPIC_API_KEY` ha priorita rispetto alle credenziali subscription/OAuth del login standard, una volta approvato il suo utilizzo.",
    )
    add_bullets(
        doc,
        [
            "Se hai gia fatto login via browser ma l'ambiente contiene `ANTHROPIC_API_KEY`, Claude Code tendera a usare la chiave API.",
            "In modalita interattiva, la prima volta ti puo chiedere se vuoi approvare l'uso della custom API key.",
            "In modalita non interattiva (`-p`), se `ANTHROPIC_API_KEY` e presente viene usata automaticamente.",
            "Se devi tornare al login standard, elimina o svuota la variabile d'ambiente e controlla poi `/status`.",
        ],
    )

    doc.add_heading("Primo avvio con API key", level=1)
    add_numbered(
        doc,
        [
            "Apri PowerShell con la variabile `ANTHROPIC_API_KEY` gia caricata.",
            "Lancia `claude`.",
            "Se Claude Code chiede di confermare l'uso della custom API key, approvala.",
            "Controlla `/status` per verificare quale metodo di autenticazione risulta attivo.",
            "Se il tuo team ti ha chiesto un comportamento diverso, puoi controllare la sezione `/config` relativa all'uso della custom API key.",
        ],
    )
    add_code_block(doc, ["claude", "/status"])

    doc.add_heading("Errori frequenti con API key", level=1)
    add_table(
        doc,
        ["Problema", "Causa probabile", "Cosa fare"],
        [
            [
                "403 `Request not allowed`",
                "Il tuo account Console non ha il ruolo corretto oppure la chiave non e autorizzata.",
                "Chiedi all'admin di verificare il ruolo `Claude Code` o `Developer` in Anthropic Console.",
            ],
            [
                "Hai una subscription attiva ma compare `This organization has been disabled`",
                "Una vecchia `ANTHROPIC_API_KEY` sta sovrascrivendo il login subscription.",
                "Rimuovi la variabile dall'ambiente corrente, dal profilo PowerShell e dalle variabili utente, poi verifica con `/status`.",
            ],
            [
                "Claude continua a usare credenziali diverse da quelle attese",
                "Ci sono piu metodi di autenticazione in gioco.",
                "Ricontrolla la presenza di `ANTHROPIC_API_KEY`, eventuali token gia salvati e l'output di `/status`.",
            ],
            [
                "La chiave sembra impostata ma non viene letta",
                "PowerShell e stato aperto prima della variabile persistente, oppure la variabile e stata salvata altrove.",
                "Chiudi e riapri PowerShell. Controlla la variabile sia nella sessione corrente sia nello scope utente.",
            ],
            [
                "Vuoi tornare al login browser standard",
                "La chiave API rimane nell'ambiente e continua ad avere priorita.",
                "Rimuovi `ANTHROPIC_API_KEY` dalla sessione, da `$PROFILE` e dalle variabili utente; poi rilancia `claude`.",
            ],
        ],
        [1.9, 1.8, 2.5],
    )

    doc.add_heading("Dove controllare su Windows se la variabile e gia impostata", level=1)
    add_bullets(
        doc,
        [
            "Sessione PowerShell corrente: `Get-ChildItem Env:ANTHROPIC_API_KEY`.",
            "Variabile utente persistente: `[Environment]::GetEnvironmentVariable(\"ANTHROPIC_API_KEY\",\"User\")`.",
            "Profilo PowerShell: apri il file indicato da `$PROFILE` e cerca eventuali righe che impostano `ANTHROPIC_API_KEY`.",
            "Interfaccia grafica: Pannello di controllo o Impostazioni avanzate di sistema -> Variabili d'ambiente.",
        ],
    )

    doc.add_heading("Quando conviene usare questa modalita", level=1)
    add_bullets(
        doc,
        [
            "Quando il tuo team usa billing API e ti ha distribuito una chiave dedicata.",
            "Quando devi standardizzare l'autenticazione da terminale senza passare dal browser su ogni macchina.",
            "Quando vuoi separare nettamente il percorso CLI basato su API dal tuo eventuale abbonamento personale Claude.ai.",
        ],
    )
    add_callout(
        doc,
        "Differenza pratica",
        "Con login Pro, Max, Team o Enterprise l'accesso parte dal browser e usa credenziali OAuth. Con API key, l'identita operativa del terminale dipende prima di tutto da `ANTHROPIC_API_KEY` e dalla configurazione del tuo ambiente.",
    )

    doc.add_heading("Riepilogo operativo", level=1)
    add_numbered(
        doc,
        [
            "Installa Claude Code con PowerShell se non e gia presente.",
            "Imposta `ANTHROPIC_API_KEY` nella sessione corrente o come variabile utente.",
            "Verifica la presenza della variabile senza stampare la chiave completa.",
            "Avvia `claude`, approva l'uso della custom API key se richiesto e controlla `/status`.",
        ],
    )

    add_sources_section(doc)
    return doc


def main():
    outputs = {
        "Guida installazione Claude Code su Windows.docx": build_standard_guide(),
        "Guida Claude Code su Windows con API key.docx": build_api_key_guide(),
    }
    for filename, document in outputs.items():
        document.save(ROOT / filename)
        print(f"Creato: {filename}")


if __name__ == "__main__":
    main()
